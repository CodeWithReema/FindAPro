import io
import json
import re
from django.conf import settings
from django.contrib.auth.decorators import login_required
from django.http import HttpResponse, JsonResponse
from django.shortcuts import redirect, get_object_or_404
from django.views.decorators.http import require_POST
from django.views.generic import TemplateView
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

from .models import BcubedResult


PRESENTATION_FORMAT_INSTRUCTIONS = (
    "Return your answer in clean, presentation-ready Markdown.\n"
    "Rules:\n"
    "- Use short headings (##) and subheadings (###) where helpful.\n"
    "- Prefer bullet points and numbered steps.\n"
    "- Keep lines concise; avoid long paragraphs.\n"
    "- When you give options, format as a list of labeled options.\n"
    "- Do NOT include disclaimers unless explicitly requested.\n"
    "- Do NOT wrap the whole response in a code block.\n"
)


class ProviderRequiredMixin:
    def dispatch(self, request, *args, **kwargs):
        if not request.user.is_authenticated:
            return redirect('accounts:login')
        if not getattr(request.user, 'has_provider_profile', False):
            return redirect('providers:join_as_pro')
        return super().dispatch(request, *args, **kwargs)


class HubView(ProviderRequiredMixin, TemplateView):
    template_name = 'bcubed/hub.html'

    def get_context_data(self, **kwargs):
        ctx = super().get_context_data(**kwargs)
        ctx['total_results'] = BcubedResult.objects.filter(user=self.request.user).count()
        return ctx


class BuildView(ProviderRequiredMixin, TemplateView):
    template_name = 'bcubed/build.html'

    def get_context_data(self, **kwargs):
        ctx = super().get_context_data(**kwargs)
        ctx['history'] = BcubedResult.objects.filter(user=self.request.user, stage='build')[:20]
        return ctx


class BrandView(ProviderRequiredMixin, TemplateView):
    template_name = 'bcubed/brand.html'

    def get_context_data(self, **kwargs):
        ctx = super().get_context_data(**kwargs)
        ctx['history'] = BcubedResult.objects.filter(user=self.request.user, stage='brand')[:20]
        return ctx


class BankView(ProviderRequiredMixin, TemplateView):
    template_name = 'bcubed/bank.html'

    def get_context_data(self, **kwargs):
        ctx = super().get_context_data(**kwargs)
        ctx['history'] = BcubedResult.objects.filter(user=self.request.user, stage='bank')[:20]
        return ctx


@login_required
@require_POST
def delete_result(request, pk):
    result = get_object_or_404(BcubedResult, pk=pk, user=request.user)
    stage = result.stage
    result.delete()
    return JsonResponse({'ok': True, 'stage': stage})


@login_required
@require_POST
def export_docx(request):
    """Export B³ result content as a Word document (.docx)."""
    data = _parse_body(request)
    content = (data.get('content') or '').strip()
    title = (data.get('title') or 'B³ Export').strip()[:200]
    if not content:
        return JsonResponse({'error': 'No content to export.'}, status=400)
    try:
        doc = Document()
        heading = doc.add_heading(title, level=0)
        heading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        for block in content.split('\n'):
            block = block.strip()
            if not block:
                doc.add_paragraph()
                continue
            p = doc.add_paragraph(block)
            p.paragraph_format.space_after = Pt(8)
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        filename = re.sub(r'[^\w\s-]', '', title)[:50].strip() or 'bcubed-export'
        filename = re.sub(r'[-\s]+', '-', filename) + '.docx'
        response = HttpResponse(buffer.getvalue(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = f'attachment; filename="{filename}"'
        return response
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=500)


def _openai_chat(system_prompt, user_prompt):
    import openai
    client = openai.OpenAI(api_key=settings.OPENAI_API_KEY)
    system_prompt = system_prompt.strip() + "\n\n" + PRESENTATION_FORMAT_INSTRUCTIONS
    response = client.chat.completions.create(
        model='gpt-4o-mini',
        messages=[
            {'role': 'system', 'content': system_prompt},
            {'role': 'user', 'content': user_prompt},
        ],
        max_tokens=800,
        temperature=0.7,
    )
    return response.choices[0].message.content.strip()


def _save_result(user, stage, tool, inputs, result):
    BcubedResult.objects.create(user=user, stage=stage, tool=tool, inputs=inputs, result=result)


def _parse_body(request):
    try:
        return json.loads(request.body)
    except (json.JSONDecodeError, ValueError):
        return {}


# ── BUILD tools ───────────────────────────────────────────────────────────────

@login_required
@require_POST
def api_pitch(request):
    data = _parse_body(request)
    business_type = data.get('business_type', '').strip()
    skills = data.get('skills', '').strip()
    target_audience = data.get('target_audience', '').strip()
    if not business_type:
        return JsonResponse({'error': 'Business type is required.'}, status=400)
    try:
        result = _openai_chat(
            system_prompt='You are an expert business coach who writes compelling elevator pitches.',
            user_prompt=(
                f'Write a compelling 3-sentence elevator pitch for a {business_type} business. '
                f'Skills: {skills}. Target audience: {target_audience}. '
                f'Make it punchy, specific, and memorable.'
            ),
        )
        _save_result(request.user, 'build', 'pitch', data, result)
        return JsonResponse({'result': result})
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=500)


@login_required
@require_POST
def api_checklist(request):
    data = _parse_body(request)
    business_type = data.get('business_type', '').strip()
    state = data.get('state', '').strip()
    if not business_type:
        return JsonResponse({'error': 'Business type is required.'}, status=400)
    try:
        result = _openai_chat(
            system_prompt='You are a business startup advisor with deep knowledge of legal and operational requirements.',
            user_prompt=(
                f'Give me a numbered startup checklist (10 items) for starting a {business_type} business'
                f'{" in " + state if state else ""}. '
                f'Be specific and practical. Include legal, financial, and operational steps.'
            ),
        )
        _save_result(request.user, 'build', 'checklist', data, result)
        return JsonResponse({'result': result})
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=500)


@login_required
@require_POST
def api_niche(request):
    data = _parse_body(request)
    skills = data.get('skills', '').strip()
    interests = data.get('interests', '').strip()
    if not skills:
        return JsonResponse({'error': 'Skills are required.'}, status=400)
    try:
        result = _openai_chat(
            system_prompt='You are a market research expert who identifies profitable business niches.',
            user_prompt=(
                f'Suggest 5 profitable market niches for someone with these skills: {skills}'
                f'{" and interests: " + interests if interests else ""}. '
                f'Format each as: Niche Name — one sentence explaining why it\'s profitable.'
            ),
        )
        _save_result(request.user, 'build', 'niche', data, result)
        return JsonResponse({'result': result})
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=500)


# ── BRAND tools ───────────────────────────────────────────────────────────────

@login_required
@require_POST
def api_names(request):
    data = _parse_body(request)
    business_type = data.get('business_type', '').strip()
    values = data.get('values', '').strip()
    style = data.get('style', '').strip()
    if not business_type:
        return JsonResponse({'error': 'Business type is required.'}, status=400)
    try:
        result = _openai_chat(
            system_prompt='You are a brand naming expert who creates memorable, distinctive business names.',
            user_prompt=(
                f'Generate 5 creative business names for a {business_type} business.'
                f'{" Values: " + values + "." if values else ""}'
                f'{" Style preference: " + style + "." if style else ""} '
                f'For each: Name — one sentence of reasoning.'
            ),
        )
        _save_result(request.user, 'brand', 'names', data, result)
        return JsonResponse({'result': result})
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=500)


@login_required
@require_POST
def api_tagline(request):
    data = _parse_body(request)
    business_name = data.get('business_name', '').strip()
    what_you_do = data.get('what_you_do', '').strip()
    if not business_name or not what_you_do:
        return JsonResponse({'error': 'Business name and description are required.'}, status=400)
    try:
        result = _openai_chat(
            system_prompt='You are a copywriter who specializes in punchy, memorable brand taglines.',
            user_prompt=(
                f'Write 5 punchy taglines for {business_name}, a {what_you_do} business. '
                f'Each tagline should be under 8 words, memorable, and highlight a key benefit.'
            ),
        )
        _save_result(request.user, 'brand', 'tagline', data, result)
        return JsonResponse({'result': result})
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=500)


@login_required
@require_POST
def api_copy(request):
    data = _parse_body(request)
    service = data.get('service', '').strip()
    target_audience = data.get('target_audience', '').strip()
    tone = data.get('tone', 'professional').strip()
    if not service:
        return JsonResponse({'error': 'Service description is required.'}, status=400)
    try:
        result = _openai_chat(
            system_prompt='You are a marketing copywriter who writes compelling service provider bios.',
            user_prompt=(
                f'Write a professional bio/description (approximately 150 words) for a {service} provider'
                f'{" targeting " + target_audience if target_audience else ""}. '
                f'Tone: {tone}. Include a strong opening, key value proposition, and a call to action.'
            ),
        )
        _save_result(request.user, 'brand', 'copy', data, result)
        return JsonResponse({'result': result})
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=500)


@login_required
@require_POST
def api_logo(request):
    data = _parse_body(request)
    business_name = data.get('business_name', '').strip()
    business_type = data.get('business_type', '').strip()
    style = data.get('style', '').strip()
    if not business_name or not business_type:
        return JsonResponse({'error': 'Business name and type are required.'}, status=400)
    try:
        result = _openai_chat(
            system_prompt='You are a professional brand designer and art director who creates detailed logo concepts.',
            user_prompt=(
                f'Create a detailed logo concept for {business_name}, a {business_type} business.'
                f'{" Style preference: " + style + "." if style else ""} '
                f'Describe: 1) Icon/symbol concept, 2) Color palette with hex codes, '
                f'3) Typography recommendation, 4) Overall feel and message. '
                f'Then briefly suggest 2 alternative concepts.'
            ),
        )
        _save_result(request.user, 'brand', 'logo', data, result)
        return JsonResponse({'result': result})
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=500)


@login_required
@require_POST
def api_seo(request):
    data = _parse_body(request)
    business_category = data.get('business_category', '').strip()
    city = data.get('city', '').strip()
    state = data.get('state', '').strip()
    if not business_category:
        return JsonResponse({'error': 'Business category is required.'}, status=400)
    location = f'{city}, {state}' if city and state else (city or state or 'your area')
    try:
        result = _openai_chat(
            system_prompt='You are an SEO specialist who helps local service businesses rank higher online.',
            user_prompt=(
                f'Give 5 actionable SEO tips for a {business_category} business in {location}. '
                f'Be specific — include keyword ideas, Google Business Profile tips, and local SEO tactics.'
            ),
        )
        _save_result(request.user, 'brand', 'seo', data, result)
        return JsonResponse({'result': result})
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=500)


# ── BANK tools ────────────────────────────────────────────────────────────────

@login_required
@require_POST
def api_pricing(request):
    data = _parse_body(request)
    service_type = data.get('service_type', '').strip()
    years_experience = data.get('years_experience', '').strip()
    city = data.get('city', '').strip()
    if not service_type:
        return JsonResponse({'error': 'Service type is required.'}, status=400)
    try:
        result = _openai_chat(
            system_prompt='You are a business pricing strategist who helps service professionals price their work competitively.',
            user_prompt=(
                f'What should a {service_type} provider'
                f'{" with " + years_experience + " years of experience" if years_experience else ""}'
                f'{" in " + city if city else ""} charge? '
                f'Give hourly and project-based pricing ranges with clear reasoning. '
                f'Include tips on when to charge more.'
            ),
        )
        _save_result(request.user, 'bank', 'pricing', data, result)
        return JsonResponse({'result': result})
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=500)


@login_required
@require_POST
def api_budget(request):
    data = _parse_body(request)
    monthly_income = data.get('monthly_income', '').strip()
    main_expenses = data.get('main_expenses', '').strip()
    if not monthly_income:
        return JsonResponse({'error': 'Monthly income is required.'}, status=400)
    try:
        result = _openai_chat(
            system_prompt='You are a financial advisor who helps freelancers and small business owners manage their money.',
            user_prompt=(
                f'Create a monthly budget plan for a freelance business with income of ${monthly_income}.'
                f'{" Main expenses: " + main_expenses + "." if main_expenses else ""} '
                f'Break it down into categories (taxes ~30%, savings, operating costs, personal), '
                f'include a savings goal, and give 2-3 money tips.'
            ),
        )
        _save_result(request.user, 'bank', 'budget', data, result)
        return JsonResponse({'result': result})
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=500)


@login_required
@require_POST
def api_tax(request):
    data = _parse_body(request)
    business_type = data.get('business_type', '').strip()
    state = data.get('state', '').strip()
    business_structure = data.get('business_structure', 'sole proprietor').strip()
    if not business_type:
        return JsonResponse({'error': 'Business type is required.'}, status=400)
    try:
        result = _openai_chat(
            system_prompt='You are a tax advisor who helps small business owners understand their tax obligations.',
            user_prompt=(
                f'What are the key tax considerations for a {business_structure} {business_type} business'
                f'{" in " + state if state else ""}? '
                f'Give 5 practical tips including deductions, quarterly payments, and record-keeping. '
                f'End with this disclaimer: "This is general information only and not legal or tax advice. '
                f'Consult a licensed CPA or tax professional for your specific situation."'
            ),
        )
        _save_result(request.user, 'bank', 'tax', data, result)
        return JsonResponse({'result': result})
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=500)


@login_required
@require_POST
def api_invoice(request):
    data = _parse_body(request)
    business_name = data.get('business_name', '').strip()
    service_name = data.get('service_name', '').strip()
    hourly_rate = data.get('hourly_rate', '').strip()
    if not business_name or not service_name:
        return JsonResponse({'error': 'Business name and service name are required.'}, status=400)
    try:
        result = _openai_chat(
            system_prompt='You are a business operations expert who creates professional invoice templates.',
            user_prompt=(
                f'Create a professional invoice template for {business_name}, a {service_name} provider'
                f'{" at $" + hourly_rate + "/hr" if hourly_rate else ""}. '
                f'Include all standard fields: invoice number, date, due date, client info, '
                f'itemized services, subtotal, tax line, total, payment terms, and thank-you note. '
                f'Format it clearly as plain text.'
            ),
        )
        _save_result(request.user, 'bank', 'invoice', data, result)
        return JsonResponse({'result': result})
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=500)
